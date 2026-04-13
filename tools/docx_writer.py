from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_resume_docx(content, filename):
    """Create a .docx file for the tailored resume."""
    doc = Document()
    # Set font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Add content as text (assuming content is structured text)
    paragraphs = content.split('\n\n')
    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para.strip())

    doc.save(filename)

def create_cover_letter_docx(content, filename, candidate_name, company_name):
    """Create a .docx file for the cover letter."""
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Header
    doc.add_paragraph(f"{candidate_name}")
    doc.add_paragraph("Your Address")
    doc.add_paragraph("City, State, ZIP")
    doc.add_paragraph("Email | Phone")
    doc.add_paragraph()
    doc.add_paragraph(f"Date")
    doc.add_paragraph()
    doc.add_paragraph(f"Hiring Manager")
    doc.add_paragraph(f"{company_name}")
    doc.add_paragraph("Company Address")
    doc.add_paragraph("City, State, ZIP")
    doc.add_paragraph()
    doc.add_paragraph("Dear Hiring Manager,")
    doc.add_paragraph()

    # Body
    paragraphs = content.split('\n\n')
    for para in paragraphs:
        doc.add_paragraph(para.strip())

    doc.save(filename)

def create_company_brief_docx(content, filename):
    """Create a .docx file for the company brief."""
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    doc.add_heading('Company Brief', level=1)
    doc.add_paragraph(content)

    doc.save(filename)